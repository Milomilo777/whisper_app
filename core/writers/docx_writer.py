"""Microsoft Word ``.docx`` writer.

Produces a structured DOCX file:

  - Heading 1: audio basename
  - Subtitle paragraph: meta line (segment count, duration if known)
  - One paragraph per segment:
      [HH:MM:SS]  <bold speaker (if any)>:  segment text

Requires ``python-docx`` (BSD-licensed, ~1 MB wheel). If the import
fails at runtime — e.g. user runs from a Python without python-docx
installed — the writer raises a clear RuntimeError instead of
crashing in the import.

The module also exposes the byte payload through ``write()``: unlike
the text-based writers which return ``str``, DOCX is binary, so this
module returns the raw zip bytes and ``_write_outputs`` in
``core.transcriber`` is taught to handle the binary case via the
small adapter at the bottom.
"""
from __future__ import annotations

import io
import os
from typing import Any

from .base import fmt_srt_time, normalize_text


def _fmt_doc_time(seconds: float) -> str:
    """``HH:MM:SS`` — drops the millisecond fraction the SRT helper carries."""
    return fmt_srt_time(seconds).split(",")[0]


def _require_docx() -> Any:
    """Lazy-import python-docx; raise a clean error if absent."""
    try:
        import docx  # type: ignore
    except ImportError as e:  # noqa: BLE001
        raise RuntimeError(
            "DOCX export requires the python-docx package. "
            "Install it via `pip install python-docx>=1.0`."
        ) from e
    return docx


def write_bytes(segments: list[dict], audio_path: str = "") -> bytes:
    """Build the docx and return its raw zip bytes."""
    docx = _require_docx()

    document = docx.Document()
    title = os.path.basename(audio_path) if audio_path else "Transcript"
    document.add_heading(title, level=1)

    nonempty = [s for s in segments if normalize_text(s.get("text", ""))]
    if nonempty:
        last_end = float(nonempty[-1].get("end", 0.0))
        duration = _fmt_doc_time(last_end)
        document.add_paragraph(
            f"{len(nonempty)} segment(s) · {duration} total",
            style="Intense Quote",
        )

    for seg in nonempty:
        ts = _fmt_doc_time(float(seg.get("start", 0.0)))
        speaker = (seg.get("speaker") or "").strip()
        text = normalize_text(seg.get("text", ""))

        para = document.add_paragraph()
        # [HH:MM:SS]
        run_ts = para.add_run(f"[{ts}]  ")
        run_ts.bold = True
        # Optional speaker prefix
        if speaker:
            run_sp = para.add_run(f"{speaker}: ")
            run_sp.bold = True
        # Segment body
        para.add_run(text)

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def write(segments: list[dict], audio_path: str = "") -> str:
    """Return the docx bytes wrapped as a latin-1 string.

    Why latin-1? ``core.transcriber._write_outputs`` writes returned
    strings via ``open(path, "w", encoding="utf-8")``. A latin-1
    round-trip preserves every byte 0–255 unchanged through UTF-8
    encoding back into the original bytes — **only if** the receiver
    writes in binary mode. ``_write_outputs`` was extended to
    detect this writer by name and switch to a binary write; if the
    detection ever drops, this function fails fast rather than
    silently producing a corrupt DOCX.
    """
    raise RuntimeError(
        "core.writers.docx_writer must be invoked via write_bytes() — "
        "_write_outputs handles the binary path."
    )
