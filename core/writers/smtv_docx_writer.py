"""SMTV transcription ``.docx`` writer.

Fills the transcription team's exact Word template (one 4-column table)
instead of building a table from scratch, so the team's borders, fonts,
and cell styling are preserved byte-for-byte. The template ships as a
bundled asset at ``core/writers/templates/smtv_template.docx``.

Template layout (1-based rows / columns, verified):

  * Row 1: a single merged title cell with two paragraphs::

        (work title)
        Transcription in (Foreign Language) - Translation in English

    ``(work title)`` is replaced with the source-file stem and
    ``(Foreign Language)`` with the detected language name. The dash is
    a real en-dash (U+2013) in the template.
  * Row 2: header row -- col1 row-number header, col2 "Time Code",
    col3 "Foreign Language" (replaced with the detected language name
    when a language was detected), col4 "English Translation".
  * Row 3 onward: transcription rows. Row 3 col3 carries a
    ``[(Foreign Language) starts]`` marker that is also language-filled.
    The template provides 31 usable rows (rows 3-33); when there are
    more segments than rows, extra rows are cloned so nothing is lost.

Per transcription row:

  * col1 = incrementing row number (1, 2, 3, ...).
  * col2 = Time Code ``HH:MM:SS.m`` (one-digit tenths-of-second) of the
    SEGMENT (phrase) start time.
  * col3 = the transcribed segment text (sanitised; speaker label
    prepended when present, matching the docx writer convention).
  * col4 = English Translation -- left EMPTY for a human translator.

Unlike the other writers (whose ``write(segments, audio_path) -> str``
contract is frozen), this writer needs the detected language and the
work title, so ``core.transcriber._write_outputs`` special-cases it and
calls :func:`write_bytes` with the extra keyword arguments. It is
registered under ``BINARY_WRITERS`` via :func:`write` which raises, so
``is_binary`` routes callers to the binary path.
"""
from __future__ import annotations

import copy
import datetime
import io
import math
import os
from typing import Any

from .base import normalize_text, sanitize_for_xml, speaker_prefix

# Placeholder strings exactly as they appear (consolidated) in the
# template's paragraph text. Replacement is run-aware so styling on the
# surrounding text survives.
_WORK_TITLE_PLACEHOLDER = "(work title)"
_FOREIGN_PLACEHOLDER = "(Foreign Language)"
# The row-2 (header) col3 literal text -- NOT parenthesised like the
# title/marker placeholders above, so it needs its own needle.
_HEADER_FOREIGN_LABEL = "Foreign Language"
# Neutral fill for the "(Foreign Language)" placeholder when no language was
# detected, so the row-0 "[... starts]" cue marker is still meaningful (and is
# never clobbered) instead of leaving the literal placeholder dangling.
_NEUTRAL_LANG_LABEL = "transcription"

# First transcription row is the 3rd table row (index 2); rows 0/1 are
# the title and header rows.
_FIRST_DATA_ROW = 2

# ISO-639-1 (and a few common variants) -> human language name. Falls
# back to the raw code when unmapped. Kept core-side and Tk-free.
_LANGUAGE_NAMES: dict[str, str] = {
    "af": "Afrikaans",
    "sq": "Albanian",
    "am": "Amharic",
    "ar": "Arabic",
    "hy": "Armenian",
    "az": "Azerbaijani",
    "eu": "Basque",
    "be": "Belarusian",
    "bn": "Bengali",
    "bs": "Bosnian",
    "bg": "Bulgarian",
    "ca": "Catalan",
    "ceb": "Cebuano",
    "zh": "Chinese (simplified)",
    "zh-CN": "Chinese (simplified)",
    "zh-TW": "Chinese (traditional)",
    "co": "Corsican",
    "hr": "Croatian",
    "cs": "Czech",
    "da": "Danish",
    "nl": "Dutch",
    "en": "English",
    "eo": "Esperanto",
    "et": "Estonian",
    "fi": "Finnish",
    "fr": "French",
    "fy": "Frisian",
    "gl": "Galician",
    "ka": "Georgian",
    "de": "German",
    "el": "Greek",
    "gu": "Gujarati",
    "ht": "Haitian Creole",
    "ha": "Hausa",
    "haw": "Hawaiian",
    "iw": "Hebrew",
    "hi": "Hindi",
    "hmn": "Hmong",
    "hu": "Hungarian",
    "is": "Icelandic",
    "ig": "Igbo",
    "id": "Indonesian",
    "ga": "Irish",
    "it": "Italian",
    "ja": "Japanese",
    "jv": "Javanese",
    "kn": "Kannada",
    "kk": "Kazakh",
    "km": "Khmer",
    "ko": "Korean",
    "ku": "Kurdish",
    "ky": "Kyrgyz",
    "lo": "Lao",
    "la": "Latin",
    "lv": "Latvian",
    "lt": "Lithuanian",
    "lb": "Luxembourgish",
    "mk": "Macedonian",
    "mg": "Malagasy",
    "ms": "Malay",
    "ml": "Malayalam",
    "mt": "Maltese",
    "mi": "Maori",
    "mr": "Marathi",
    "mn": "Mongolian",
    "my": "Myanmar (Burmese)",
    "ne": "Nepali",
    "no": "Norwegian",
    "ny": "Nyanja (Chichewa)",
    "ps": "Pashto",
    "fa": "Persian",
    "pl": "Polish",
    "pt": "Portuguese",
    "pa": "Punjabi",
    "ro": "Romanian",
    "ru": "Russian",
    "sm": "Samoan",
    "gd": "Scots Gaelic",
    "sr": "Serbian",
    "st": "Sesotho",
    "sn": "Shona",
    "sd": "Sindhi",
    "si": "Sinhala (Sinhalese)",
    "sk": "Slovak",
    "sl": "Slovenian",
    "so": "Somali",
    "es": "Spanish",
    "su": "Sundanese",
    "sw": "Swahili",
    "sv": "Swedish",
    "tl": "Tagalog (Filipino)",
    "tg": "Tajik",
    "ta": "Tamil",
    "te": "Telugu",
    "th": "Thai",
    "tr": "Turkish",
    "uk": "Ukrainian",
    "ur": "Urdu",
    "uz": "Uzbek",
    "vi": "Vietnamese",
    "cy": "Welsh",
    "xh": "Xhosa",
    "yi": "Yiddish",
    "yo": "Yoruba",
    "zu": "Zulu",
}


def language_name(code: str) -> str:
    """Human language name for an ISO code; the raw code if unmapped.

    Region/script suffixes are stripped (``pt-BR`` -> ``pt``) before the
    lookup so a BCP-47 code still resolves.
    """
    raw = (code or "").strip()
    if not raw:
        return ""
    key = raw.replace("_", "-").split("-")[0].lower()
    return _LANGUAGE_NAMES.get(key, raw)


def _fmt_smtv_time(seconds: float) -> str:
    """``HH:MM:SS.m`` -- 2-digit hours, MM/SS < 60, one-digit tenths.

    Non-finite / negative inputs clamp to ``00:00:00.0`` so a buggy
    backend timestamp never produces a garbage cell.
    """
    if seconds is None or not isinstance(seconds, (int, float)):
        seconds = 0.0
    seconds = float(seconds)
    if not math.isfinite(seconds) or seconds < 0:
        seconds = 0.0
    # Round to tenths first so 59.96 doesn't render as 60.0 in the
    # seconds field (it should roll into the next minute). Use round-HALF-UP
    # (floor(x + 0.5)) rather than Python's built-in round(), which is
    # banker's rounding (round-half-to-even): round(0.05*10)=0 but
    # round(0.15*10)=2 makes equal .x5 inputs round inconsistently, so
    # timecodes were not predictable. floor(x+0.5) always rounds .x5 up.
    total_tenths = int(math.floor(seconds * 10 + 0.5))
    tenths = total_tenths % 10
    total_secs = total_tenths // 10
    hours, rem = divmod(total_secs, 3600)
    minutes, sec = divmod(rem, 60)
    return f"{hours:02d}:{minutes:02d}:{sec:02d}.{tenths}"


def _require_docx() -> Any:
    """Lazy-import python-docx; raise a clean error if absent."""
    try:
        import docx  # type: ignore
    except ImportError as e:  # noqa: BLE001
        raise RuntimeError(
            "SMTV transcription export requires the python-docx package. "
            "Install it via `pip install python-docx>=1.0`."
        ) from e
    return docx


def template_path() -> str:
    """Absolute path to the bundled SMTV template.

    Resolves via ``core.paths.resource_base`` (works in source, onefile,
    onedir, and the embed tree). Falls back to the path next to this
    module if the bundled copy is missing.
    """
    from core.paths import resource_base

    bundled = os.path.join(
        resource_base(), "core", "writers", "templates", "smtv_template.docx"
    )
    if os.path.isfile(bundled):
        return bundled
    # Fallback: alongside this source file (covers an odd resource_base
    # in some test/frozen layouts).
    local = os.path.join(
        os.path.dirname(os.path.abspath(__file__)), "templates", "smtv_template.docx"
    )
    return local


def _replace_in_paragraph(paragraph: Any, needle: str, replacement: str) -> None:
    """Replace ``needle`` with ``replacement`` across a paragraph.

    python-docx splits text into arbitrary runs, so the placeholder can
    straddle several runs. We consolidate the paragraph's full text,
    substitute, then write the result back into the FIRST run (keeping
    its formatting) and blank the remaining runs. This collapses per-run
    styling differences inside the span, which is fine for these
    all-similarly-formatted placeholders and leaves the rest of the
    document untouched.
    """
    runs = paragraph.runs
    if not runs:
        return
    full = "".join(r.text for r in runs)
    if needle not in full:
        return
    new = full.replace(needle, replacement)
    runs[0].text = new
    for r in runs[1:]:
        r.text = ""


def _clear_cell_text(cell: Any) -> None:
    """Empty a table cell's text while keeping one (formatted) run."""
    for para in cell.paragraphs:
        runs = para.runs
        if runs:
            runs[0].text = ""
            for r in runs[1:]:
                r.text = ""


def _set_cell_text(cell: Any, text: str) -> None:
    """Set a cell's text into its first run (preserving cell styling).

    The template's data cells start empty; we write into the first run
    of the first paragraph so the team's cell font/borders are kept. If
    a cell has no run yet, add one.
    """
    paras = cell.paragraphs
    if not paras:
        cell.add_paragraph(text)
        return
    para = paras[0]
    if not para.runs:
        para.add_run(text)
    else:
        para.runs[0].text = text
        for r in para.runs[1:]:
            r.text = ""
    # Blank any trailing paragraphs in the cell.
    for extra in paras[1:]:
        for r in extra.runs:
            r.text = ""


def _clone_last_row(table: Any) -> Any:
    """Append a deep copy of the last table row and return it.

    Cloning the existing row's XML preserves the team's row height, cell
    widths, borders, and fonts. The clone's cell text is cleared by the
    caller before filling.
    """
    last_tr = table.rows[-1]._tr
    new_tr = copy.deepcopy(last_tr)
    last_tr.addnext(new_tr)
    return table.rows[-1]


def write_bytes(
    segments: list[dict],
    audio_path: str = "",
    *,
    language: str = "",
    work_title: str = "",
) -> bytes:
    """Fill the SMTV template and return the resulting docx bytes.

    ``language`` is the detected ISO code (mapped to a human name);
    ``work_title`` is the source-file stem. When either is empty the
    corresponding template placeholder is left verbatim.
    """
    _require_docx()
    import docx  # type: ignore

    document = docx.Document(template_path())

    lang_label = language_name(language)
    title = sanitize_for_xml((work_title or "").strip())

    if not document.tables:
        raise RuntimeError("SMTV template is malformed: no table found.")
    table = document.tables[0]

    # --- Row 1: title cell placeholders -------------------------------
    title_cell = table.rows[0].cells[0]
    for para in title_cell.paragraphs:
        if title:
            _replace_in_paragraph(para, _WORK_TITLE_PLACEHOLDER, title)
        if lang_label:
            _replace_in_paragraph(para, _FOREIGN_PLACEHOLDER, lang_label)

    # --- Row 3 col3: "[(Foreign Language) starts]" marker -------------
    # Fill the cue even when no language was detected: use a neutral label so
    # the "[... starts]" marker stays meaningful and is preserved below (the
    # append branch keys off this same label, so the cue is never clobbered).
    marker_label = lang_label or _NEUTRAL_LANG_LABEL
    if len(table.rows) > _FIRST_DATA_ROW:
        marker_cell = table.rows[_FIRST_DATA_ROW].cells[2]
        for para in marker_cell.paragraphs:
            _replace_in_paragraph(para, _FOREIGN_PLACEHOLDER, marker_label)

    # --- Row 2 col3 header: "Foreign Language" -> detected language ---
    # Only replace when a language was actually detected; with none, the
    # generic header stays so the template still reads sensibly.
    if lang_label and len(table.rows) > 1:
        header_cell = table.rows[1].cells[2]
        for para in header_cell.paragraphs:
            _replace_in_paragraph(para, _HEADER_FOREIGN_LABEL, lang_label)

    # --- Transcription rows -------------------------------------------
    nonempty = [s for s in segments if normalize_text(str(s.get("text") or ""))]

    for idx, seg in enumerate(nonempty):
        row_index = _FIRST_DATA_ROW + idx
        if row_index >= len(table.rows):
            new_row = _clone_last_row(table)
            for c in new_row.cells:
                _clear_cell_text(c)
        row = table.rows[row_index]
        cells = row.cells

        row_number = str(idx + 1)
        time_code = _fmt_smtv_time(float(seg.get("start", 0.0) or 0.0))
        prefix = speaker_prefix(seg)
        body = sanitize_for_xml(prefix + normalize_text(str(seg.get("text") or "")))

        # col1 = row number
        _set_cell_text(cells[0], row_number)
        # col2 = time code
        _set_cell_text(cells[1], time_code)
        # col3 = foreign-language text. The first data row already had the
        # "[<Lang> starts]" marker filled above (always, even with no detected
        # language); APPEND the segment text after it (in a new run) so the
        # team's cue is preserved instead of being overwritten.
        if idx == 0:
            marker_para = cells[2].paragraphs[0]
            marker_para.add_run(" " + body)
        else:
            _set_cell_text(cells[2], body)
        # col4 = English Translation -> left empty for a translator.
        _set_cell_text(cells[3], "")

    # The template's own "modified" timestamp would otherwise carry straight
    # through into every generated docx, making each output look like it was
    # last touched whenever the template was authored rather than now.
    document.core_properties.modified = datetime.datetime.now(datetime.timezone.utc)

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def write(segments: list[dict], audio_path: str = "") -> str:
    """Always raises -- SMTV docx is binary; use :func:`write_bytes`.

    Registered in ``BINARY_WRITERS`` so ``is_binary`` routes callers to
    the binary path. If that detection ever drops this fails fast rather
    than silently producing a corrupt file.
    """
    raise RuntimeError(
        "core.writers.smtv_docx_writer must be invoked via write_bytes() — "
        "_write_outputs handles the binary path with language/work_title."
    )
