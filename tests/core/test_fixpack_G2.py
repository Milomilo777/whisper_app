"""Fixpack G2 regression tests for ``core.writers.smtv_docx_writer``.

Covers the data-loss bug where a segment whose ``text`` key is present
but explicitly ``None`` was stringified to the literal word ``"None"``
and written into the transcription cell (instead of being treated as
empty and dropped). Resumed-checkpoint JSON and hand-edited JSON can
legitimately carry ``text: null``.

Hermetic: no Tk root, no network, no real model/binaries -- only
``python-docx`` parsing of the in-memory bytes the writer returns.
"""
from __future__ import annotations

import io as _io

import pytest

pytest.importorskip("docx")

from core.writers import smtv_docx_writer


def _smtv_table(payload: bytes):
    from docx import Document  # type: ignore

    document = Document(_io.BytesIO(payload))
    assert document.tables, "SMTV docx should contain one table"
    return document.tables[0]


def test_smtv_text_null_segment_is_dropped_not_rendered_as_None():
    """A segment with ``text: None`` must be filtered out entirely, so it
    contributes no row and never injects the literal string ``"None"``."""
    segs = [
        {"start": 1.0, "end": 2.0, "text": None},
        {"start": 2.0, "end": 3.0, "text": "real line"},
    ]
    payload = smtv_docx_writer.write_bytes(
        segs, "t.mp4", language="fa", work_title="t"
    )
    table = _smtv_table(payload)

    # The literal word "None" must not appear anywhere in the document.
    full_text = "\n".join(
        cell.text for row in table.rows for cell in row.cells
    )
    assert "None" not in full_text

    # Only the one real segment survives: it lands on the first data row
    # (row index 2) appended after the "[Persian starts]" marker, numbered
    # "1". There must be no second numbered data row.
    first = [table.rows[2].cells[i].text for i in range(4)]
    assert first[0] == "1"
    assert "[Persian starts]" in first[2]
    assert "real line" in first[2]

    # No further numbered transcription row exists (the null segment did
    # not create one).
    if len(table.rows) > 3:
        assert table.rows[3].cells[0].text == ""


def test_smtv_all_text_null_yields_no_data_rows():
    """When every segment's text is None, the marker row stays bare (no
    appended body) and no numbered transcription rows are produced."""
    segs = [
        {"start": 0.0, "end": 1.0, "text": None},
        {"start": 1.0, "end": 2.0, "text": None},
    ]
    payload = smtv_docx_writer.write_bytes(
        segs, "t.mp4", language="ko", work_title="t"
    )
    table = _smtv_table(payload)

    full_text = "\n".join(
        cell.text for row in table.rows for cell in row.cells
    )
    assert "None" not in full_text

    # The marker cell still holds the cue, but no segment body was appended
    # and no row was numbered "1".
    first = [table.rows[2].cells[i].text for i in range(4)]
    assert first[0] == ""
    assert "[Korean starts]" in first[2]


def test_smtv_nonstring_text_still_tolerated():
    """The original ``str(...)`` wrapper guarded against non-string text
    (e.g. an int from hand-edited JSON); the fix must keep that working."""
    segs = [{"start": 0.0, "end": 1.0, "text": 123}]
    payload = smtv_docx_writer.write_bytes(
        segs, "t.mp4", language="en", work_title="t"
    )
    table = _smtv_table(payload)
    first = [table.rows[2].cells[i].text for i in range(4)]
    assert first[0] == "1"
    assert "123" in first[2]
