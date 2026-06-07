"""Tests for ``core.writers`` — every writer takes the same fixture."""
from __future__ import annotations

import json
import os

import pytest

from core.writers import (
    BINARY_WRITERS,
    WRITERS,
    docx_writer,
    get_binary_writer,
    get_writer,
    is_binary,
    json_writer,
    lrc,
    md,
    srt,
    supported_formats,
    tsv,
    txt,
    vtt,
)
from core.writers.base import (
    escape_cue_separator,
    fmt_lrc_time,
    fmt_srt_time,
    fmt_vtt_time,
    normalize_text,
    sanitize_for_xml,
    speaker_prefix,
)


@pytest.fixture
def segments():
    return [
        {"start": 0.0, "end": 1.5, "text": "Hello world"},
        {
            "start": 1.5,
            "end": 3.25,
            "text": "Second   line",
            "words": [
                {"start": 1.5, "end": 2.0, "word": "Second", "probability": 0.9},
                {"start": 2.0, "end": 3.25, "word": "line", "probability": 0.8},
            ],
        },
        {"start": 3.25, "end": 5.0, "text": "Third\tline\nwith\twhitespace"},
    ]


# --- base helpers ---------------------------------------------------------


def test_fmt_srt_time_zero():
    assert fmt_srt_time(0) == "00:00:00,000"


def test_fmt_srt_time_handles_hours():
    assert fmt_srt_time(3661.5) == "01:01:01,500"


def test_fmt_srt_time_negative_clamped():
    assert fmt_srt_time(-0.1) == "00:00:00,000"


def test_fmt_vtt_time_uses_period():
    assert fmt_vtt_time(1.234) == "00:00:01.234"


def test_fmt_lrc_time_format():
    assert fmt_lrc_time(75.5) == "[01:15.50]"


def test_normalize_text_collapses_internal_whitespace():
    assert normalize_text(" hello\t  world\n") == "hello world"


def test_normalize_text_handles_none():
    assert normalize_text("") == ""


# --- writers --------------------------------------------------------------


def test_srt_writer_basic_shape(segments):
    body = srt.write(segments)
    assert body.startswith("1\n")
    assert "00:00:00,000 --> 00:00:01,500" in body
    assert "Hello world" in body
    assert "00:00:03,250 --> 00:00:05,000" in body


def test_srt_writer_segment_count(segments):
    body = srt.write(segments)
    assert body.count(" --> ") == len(segments)


def test_vtt_writer_starts_with_header(segments):
    body = vtt.write(segments)
    assert body.startswith("WEBVTT\n")
    assert " --> " in body


def test_vtt_writer_emits_karaoke_when_words_present(segments):
    body = vtt.write(segments)
    assert "<00:00:01.500><c>Second</c>" in body


def test_tsv_writer_has_header_row(segments):
    body = tsv.write(segments)
    assert body.startswith("start\tend\ttext\n")
    rows = body.strip().split("\n")
    assert len(rows) == 1 + len(segments)


def test_tsv_writer_strips_internal_tabs(segments):
    body = tsv.write(segments)
    data_lines = body.strip().split("\n")[1:]
    for line in data_lines:
        cells = line.split("\t")
        assert len(cells) == 3
        assert "\t" not in cells[2]


def test_txt_writer_one_segment_per_line(segments):
    body = txt.write(segments)
    lines = body.strip().split("\n")
    assert len(lines) == len(segments)
    assert lines[0] == "Hello world"


def test_json_writer_returns_valid_json_list(segments):
    body = json_writer.write(segments)
    parsed = json.loads(body)
    assert isinstance(parsed, list)
    assert len(parsed) == len(segments)
    assert parsed[0]["text"] == "Hello world"


def test_json_writer_preserves_words_when_present(segments):
    parsed = json.loads(json_writer.write(segments))
    assert "words" in parsed[1]
    assert parsed[1]["words"][0]["word"] == "Second"
    assert parsed[1]["words"][0]["probability"] == 0.9


def test_json_writer_omits_words_when_absent(segments):
    parsed = json.loads(json_writer.write(segments))
    assert "words" not in parsed[0]


def test_lrc_writer_uses_audio_file_name(segments):
    body = lrc.write(segments, "/tmp/song.mp3")
    assert body.startswith("[ti:song]")
    assert "[00:00.00]" in body


def test_lrc_writer_no_title_when_path_empty(segments):
    body = lrc.write(segments)
    assert not body.startswith("[ti:")


# --- registry -------------------------------------------------------------


def test_get_writer_returns_callable_for_each_format():
    for name in supported_formats():
        fn = get_binary_writer(name) if is_binary(name) else get_writer(name)
        assert callable(fn)


def test_get_writer_is_case_insensitive():
    assert get_writer("SRT") is srt.write
    assert get_writer("Vtt") is vtt.write


def test_get_writer_raises_for_unknown_format():
    with pytest.raises(KeyError):
        get_writer("xml")


def test_supported_formats_includes_canonical_set():
    formats = supported_formats()
    for required in ("srt", "vtt", "tsv", "txt", "json", "lrc", "md", "docx", "pdf"):
        assert required in formats


def test_pdf_writer_returns_pdf_magic(segments):
    from core.writers import pdf_writer
    payload = pdf_writer.write_bytes(segments, "interview.mp4")
    assert isinstance(payload, bytes)
    # PDF magic is "%PDF-".
    assert payload[:5] == b"%PDF-", payload[:8]
    # Body should be more than a few bytes — reportlab's smallest
    # output is ~ 1 KB even for one paragraph.
    assert len(payload) > 1000


def test_pdf_writer_handles_empty_and_speakers(tmp_path):
    from core.writers import pdf_writer
    # Empty segments still produces a valid PDF skeleton.
    payload = pdf_writer.write_bytes([], "")
    assert payload[:5] == b"%PDF-"

    enriched = [
        {"start": 0.0, "end": 1.0, "text": "hi", "speaker": "Speaker 00"},
        {"start": 1.0, "end": 2.0, "text": "yo"},
    ]
    payload2 = pdf_writer.write_bytes(enriched, "x.wav")
    pdf_path = tmp_path / "x.pdf"
    pdf_path.write_bytes(payload2)
    # Spot-check by writing it to disk and reading the raw bytes for
    # the segment strings. PDF content streams are compressed by
    # default, so we just verify the file is a PDF and accept that
    # deep content check belongs in an integration test.
    assert pdf_path.stat().st_size > 1000


def test_writers_handle_empty_segment_list():
    for name in supported_formats():
        # smtv_docx's registry adapter deliberately raises (it must be
        # driven via write_bytes(..., language=, work_title=) by
        # _write_outputs); its empty-input behaviour is covered by its
        # own test below.
        if name == "smtv_docx":
            continue
        if is_binary(name):
            payload = get_binary_writer(name)([], "")
            assert isinstance(payload, (bytes, bytearray))
        else:
            body = get_writer(name)([], "")
            assert isinstance(body, str)


def test_srt_writer_uses_comma_decimal():
    body = srt.write([{"start": 0.5, "end": 1.0, "text": "x"}])
    assert "00:00:00,500" in body
    assert "00:00:00.500" not in body


def test_srt_writer_prepends_speaker_when_present():
    body = srt.write([
        {"start": 0.0, "end": 1.0, "text": "alpha", "speaker": "Speaker 00"},
        {"start": 1.0, "end": 2.0, "text": "beta"},  # no speaker
    ])
    assert "Speaker 00: alpha" in body
    # The unlabeled segment stays clean — no stray "Speaker" prefix.
    lines = body.splitlines()
    beta_line = next(ln for ln in lines if ln == "beta")
    assert beta_line == "beta"


def test_json_writer_includes_speaker_when_present():
    body = json_writer.write([
        {"start": 0.0, "end": 1.0, "text": "alpha", "speaker": "Speaker 00"},
        {"start": 1.0, "end": 2.0, "text": "beta"},
    ])
    payload = json.loads(body)
    assert payload[0]["speaker"] == "Speaker 00"
    assert "speaker" not in payload[1]


# --- Markdown writer ------------------------------------------------------


def test_md_writer_has_heading_and_timestamps(segments):
    body = md.write(segments, "interview.mp3")
    assert body.startswith("# interview.mp3")
    assert "**00:00:00**" in body
    assert "**00:00:01**" in body  # second segment starts at 1.5s
    assert "Hello world" in body
    assert "Second line" in body


def test_md_writer_includes_speaker_label_when_present():
    body = md.write([
        {"start": 0.0, "end": 1.0, "text": "hi", "speaker": "Speaker 1"},
        {"start": 1.0, "end": 2.0, "text": "yo", "speaker": "Speaker 2"},
    ], "x.wav")
    assert "_Speaker 1:_" in body
    assert "_Speaker 2:_" in body


def test_md_writer_omits_speaker_when_absent():
    body = md.write([{"start": 0.0, "end": 1.0, "text": "hi"}], "x.wav")
    assert "_" not in body.replace("\n", "")  # no italic markers


def test_md_writer_skips_empty_segments():
    body = md.write([
        {"start": 0.0, "end": 1.0, "text": "  "},
        {"start": 1.0, "end": 2.0, "text": "real"},
    ], "x.wav")
    assert body.count("**") == 2  # only one segment timestamp -> one pair of **


# --- DOCX writer ----------------------------------------------------------


def test_docx_writer_returns_zip_bytes(segments):
    payload = docx_writer.write_bytes(segments, "meeting.mp4")
    assert isinstance(payload, bytes)
    # DOCX is a ZIP archive; magic bytes are "PK\x03\x04".
    assert payload[:4] == b"PK\x03\x04", payload[:8]
    # Should be substantially larger than the smallest possible zip
    # (the docx skeleton itself runs ~ 20 KB minimum).
    assert len(payload) > 5_000


def test_docx_writer_embeds_segment_text_and_speaker(segments, tmp_path):
    # Round-trip through python-docx: read the zip back and confirm
    # the segment text + speaker actually landed in document.xml.
    import zipfile

    enriched = [
        {"start": 0.0, "end": 1.0, "text": "hello", "speaker": "Alice"},
        {"start": 1.0, "end": 2.0, "text": "world", "speaker": "Bob"},
    ]
    payload = docx_writer.write_bytes(enriched, "x.wav")
    docx_path = tmp_path / "out.docx"
    docx_path.write_bytes(payload)

    with zipfile.ZipFile(docx_path) as zf:
        with zf.open("word/document.xml") as f:
            document_xml = f.read().decode("utf-8")

    assert "hello" in document_xml
    assert "world" in document_xml
    assert "Alice" in document_xml
    assert "Bob" in document_xml
    # Heading should carry the audio basename
    assert "x.wav" in document_xml
    # Timestamps appear in [HH:MM:SS] form
    assert "[00:00:00]" in document_xml
    assert "[00:00:01]" in document_xml


def test_docx_writer_handles_empty_segments_gracefully():
    payload = docx_writer.write_bytes([], "")
    assert payload[:4] == b"PK\x03\x04"


def test_is_binary_table():
    assert is_binary("docx") is True
    assert is_binary("DOCX") is True  # case-insensitive
    for name in ("srt", "vtt", "tsv", "txt", "json", "lrc", "md"):
        assert is_binary(name) is False


def test_binary_writers_registry_contains_docx_and_pdf():
    assert set(BINARY_WRITERS.keys()) == {"docx", "pdf", "smtv_docx"}


# ---------- Audit-driven safety nets ----------------------------------------


def test_speaker_prefix_handles_numeric_label():
    """A hand-edited JSON might carry a numeric speaker label — the
    writers used to crash with `int.strip() AttributeError`."""
    assert speaker_prefix({"speaker": 123}) == "123: "
    assert speaker_prefix({"speaker": None}) == ""
    assert speaker_prefix({"speaker": ""}) == ""
    assert speaker_prefix({"speaker": "  Alice  "}) == "Alice: "
    assert speaker_prefix({}) == ""


def test_sanitize_for_xml_strips_control_chars():
    """python-docx rejects XML-illegal control chars with a
    ValueError — sanitize_for_xml removes them so the DOCX writer
    can ship transcripts that contain weird tokens."""
    raw = "ok\x00\x07\x1bok"
    assert sanitize_for_xml(raw) == "okok"
    # Tab / newline / cr are XML-valid whitespace and must survive.
    assert sanitize_for_xml("a\tb\nc\rd") == "a\tb\nc\rd"


def test_escape_cue_separator_replaces_arrow_in_text():
    """Literal '-->' in segment text breaks SRT/VTT parsers that
    interpret it as a timecode line; we replace it with a unicode
    arrow that reads the same."""
    assert escape_cue_separator("foo --> bar") == "foo → bar"
    assert escape_cue_separator(None) == ""  # type: ignore[arg-type]


def test_srt_writer_speaker_int_does_not_crash():
    """Regression: writer used to crash on numeric speaker."""
    segs = [{"start": 0.0, "end": 1.0, "text": "hi", "speaker": 7}]
    out = srt.write(segs)
    assert "7: hi" in out


def test_md_writer_speaker_int_does_not_crash():
    segs = [{"start": 0.0, "end": 1.0, "text": "hi", "speaker": 42}]
    out = md.write(segs)
    assert "_42:_" in out


def test_docx_writer_strips_control_chars_from_text():
    """A NUL byte in segment text used to crash python-docx with
    `ValueError: All strings must be XML compatible`."""
    segs = [{"start": 0.0, "end": 1.0, "text": "hello\x00world"}]
    payload = docx_writer.write_bytes(segs, "x.wav")
    assert payload[:4] == b"PK\x03\x04"  # valid docx


def test_srt_writer_escapes_literal_arrow_in_text():
    """A segment whose text contains '-->' must not produce a
    parser-confusing SRT cue."""
    segs = [{"start": 0.0, "end": 1.0, "text": "code: x --> y"}]
    out = srt.write(segs)
    # Only one timecode line should contain the arrow (the cue one).
    arrow_lines = [ln for ln in out.splitlines() if "-->" in ln]
    assert len(arrow_lines) == 1


def test_json_writer_drops_nan_and_inf():
    """Strict JSON parsers reject NaN/Infinity; we clamp to safe defaults."""
    segs = [{"start": float("nan"), "end": float("inf"), "text": "hi"}]
    out = json_writer.write(segs)
    parsed = json.loads(out)
    assert parsed[0]["start"] == 0.0
    assert parsed[0]["end"] == 0.0


def test_json_writer_preserves_numeric_speaker_label():
    segs = [{"start": 0.0, "end": 1.0, "text": "hi", "speaker": 7}]
    parsed = json.loads(json_writer.write(segs))
    assert parsed[0]["speaker"] == "7"


def test_srt_fmt_clamps_nan_to_zero():
    assert fmt_srt_time(float("nan")) == "00:00:00,000"
    assert fmt_srt_time(float("inf")) == "00:00:00,000"


def test_vtt_writer_tolerates_word_start_none():
    """Regression (P2-13): a word dict with an explicit start=None (a
    hand-edited / externally-produced JSON re-fed for re-export) used to
    raise TypeError in float(None) and abort the whole VTT write. The
    word should fall back to the segment start instead of crashing."""
    segs = [{
        "start": 2.0, "end": 4.0, "text": "hi there",
        "words": [
            {"start": None, "end": 2.5, "word": "hi", "probability": 0.9},
            {"start": 2.5, "end": 4.0, "word": "there", "probability": 0.8},
        ],
    }]
    body = vtt.write(segs)  # must not raise
    assert body.startswith("WEBVTT\n")
    # The None-start word falls back to the segment start (2.0s).
    assert "<00:00:02.000><c>hi</c>" in body
    assert "<00:00:02.500><c>there</c>" in body


def test_json_writer_carries_suspect_flags():
    """The hallucination detector sets seg['suspect']/['suspect_reason'];
    the JSON writer must persist both so the transcript viewer's red-row
    feature can fire and the viewer's Save round-trips them ([19])."""
    segs = [
        {"start": 0.0, "end": 1.0, "text": "ok"},
        {"start": 1.0, "end": 2.0, "text": "you you you",
         "suspect": True, "suspect_reason": "repetition"},
    ]
    parsed = json.loads(json_writer.write(segs))
    assert "suspect" not in parsed[0]
    assert "suspect_reason" not in parsed[0]
    assert parsed[1]["suspect"] is True
    assert parsed[1]["suspect_reason"] == "repetition"


# ---------- SMTV transcription writer ---------------------------------------


def _smtv_table(payload: bytes):
    """Parse SMTV docx bytes and return its single 4-column table."""
    import io as _io

    from docx import Document  # type: ignore

    document = Document(_io.BytesIO(payload))
    assert document.tables, "SMTV docx should contain one table"
    return document.tables[0]


def test_smtv_writer_returns_valid_docx_zip():
    from core.writers import smtv_docx_writer

    segs = [{"start": 0.0, "end": 1.0, "text": "hello"}]
    payload = smtv_docx_writer.write_bytes(
        segs, "Episode 1.mp4", language="ko", work_title="Episode 1"
    )
    assert isinstance(payload, bytes)
    # DOCX is a ZIP archive; magic bytes are "PK\x03\x04".
    assert payload[:4] == b"PK\x03\x04", payload[:8]


def test_smtv_writer_fills_title_header_and_marker():
    from core.writers import smtv_docx_writer

    segs = [
        {"start": 0.0, "end": 1.0, "text": "first phrase"},
        {"start": 83.4, "end": 86.0, "text": "second", "speaker": "Speaker 1"},
    ]
    payload = smtv_docx_writer.write_bytes(
        segs, "My Show.mp4", language="ko", work_title="My Show"
    )
    table = _smtv_table(payload)

    # Title cell: source name on line 1, language on line 2 (en-dash).
    title_cell = table.rows[0].cells[0]
    assert title_cell.paragraphs[0].text == "My Show"
    line2 = title_cell.paragraphs[1].text
    assert "Transcription in Korean" in line2
    assert "– Translation in English" in line2  # real en-dash
    # Placeholders fully replaced.
    assert "(work title)" not in title_cell.text
    assert "(Foreign Language)" not in title_cell.text

    # Header row (row 2): Time Code / Foreign Language / English Translation.
    header = [table.rows[1].cells[i].text for i in range(4)]
    assert header[1] == "Time Code"
    assert header[2] == "Foreign Language"
    assert header[3] == "English Translation"

    # First data row carries the language-filled "[<Lang> starts]" cue
    # followed by the first segment text.
    first = [table.rows[2].cells[i].text for i in range(4)]
    assert first[0] == "1"
    assert first[1] == "00:00:00.0"
    assert "[Korean starts]" in first[2]
    assert "first phrase" in first[2]
    assert first[3] == ""  # English Translation left empty


def test_smtv_writer_row_numbers_timecodes_and_speaker():
    from core.writers import smtv_docx_writer

    segs = [
        {"start": 0.0, "end": 1.0, "text": "a"},
        {"start": 83.44, "end": 90.0, "text": "b", "speaker": "Speaker 2"},
        {"start": 3661.96, "end": 3663.0, "text": "c"},
    ]
    payload = smtv_docx_writer.write_bytes(
        segs, "x.mp4", language="vi", work_title="x"
    )
    table = _smtv_table(payload)

    # Incrementing row numbers.
    assert [table.rows[2 + i].cells[0].text for i in range(3)] == ["1", "2", "3"]
    # Time Code HH:MM:SS.m (one-digit tenths; rolls hours/minutes).
    assert table.rows[2].cells[1].text == "00:00:00.0"
    assert table.rows[3].cells[1].text == "00:01:23.4"
    assert table.rows[4].cells[1].text == "01:01:02.0"
    # Speaker label prepended to the foreign-language column.
    assert "Speaker 2: b" in table.rows[3].cells[2].text
    # English Translation column is empty on every data row.
    for i in range(3):
        assert table.rows[2 + i].cells[3].text == ""


def test_smtv_writer_grows_table_beyond_template_rows():
    """The template ships 31 usable rows (rows 3-33); more segments
    must clone rows so nothing is lost."""
    from core.writers import smtv_docx_writer

    n = 40  # > 31 usable template rows
    segs = [{"start": float(i), "end": float(i) + 1, "text": f"seg {i}"} for i in range(n)]
    payload = smtv_docx_writer.write_bytes(
        segs, "Big.mp4", language="ko", work_title="Big"
    )
    table = _smtv_table(payload)
    # 2 leading rows (title + header) + one row per segment.
    assert len(table.rows) == 2 + n
    # Last cloned row holds the last segment, numbered n, English empty.
    last = table.rows[-1]
    assert last.cells[0].text == str(n)
    assert "seg 39" in last.cells[2].text
    assert last.cells[3].text == ""


def test_smtv_writer_unknown_language_falls_back_to_code():
    from core.writers import smtv_docx_writer

    # An unmapped code is used verbatim as the label.
    assert smtv_docx_writer.language_name("xx") == "xx"
    assert smtv_docx_writer.language_name("pt-BR") == "Portuguese"
    assert smtv_docx_writer.language_name("") == ""


def test_smtv_time_format_clamps_and_rolls():
    from core.writers import smtv_docx_writer as w

    assert w._fmt_smtv_time(0) == "00:00:00.0"
    assert w._fmt_smtv_time(83.44) == "00:01:23.4"
    # Rounding 59.96 -> 60.0s must roll into the next minute, not 60s.
    assert w._fmt_smtv_time(59.96) == "00:01:00.0"
    # Non-finite / negative clamp to zero.
    assert w._fmt_smtv_time(float("nan")) == "00:00:00.0"
    assert w._fmt_smtv_time(-5.0) == "00:00:00.0"


def test_smtv_time_format_rounds_half_up_predictably():
    """Tenths use round-HALF-UP, not banker's rounding (round-half-to-even).

    With Python's built-in round(), round(0.05*10)=0 but round(0.15*10)=2, so
    equal ``.x5`` inputs rounded inconsistently. floor(x+0.5) always rounds
    the half up, giving predictable timecodes.
    """
    from core.writers import smtv_docx_writer as w

    assert w._fmt_smtv_time(0.05) == "00:00:00.1"   # banker's gave .0
    assert w._fmt_smtv_time(0.15) == "00:00:00.2"
    assert w._fmt_smtv_time(0.25) == "00:00:00.3"   # banker's gave .2
    assert w._fmt_smtv_time(1.25) == "00:00:01.3"   # banker's gave .2


def test_smtv_writer_empty_language_keeps_starts_cue():
    """With no detected language the row-0 "[... starts]" cue must survive.

    Regression: when language was '' the cue marker cell was OVERWRITTEN with
    just the segment body (the append branch was gated on lang_label), so the
    team lost the cue. Now a neutral label fills the placeholder and the body
    is appended after the cue.
    """
    from core.writers import smtv_docx_writer

    segs = [{"start": 0.0, "end": 1.0, "text": "first phrase"}]
    payload = smtv_docx_writer.write_bytes(
        segs, "Mystery.mp4", language="", work_title="Mystery"
    )
    table = _smtv_table(payload)
    first_marker = table.rows[2].cells[2].text
    # The "[... starts]" cue is preserved (no dangling raw placeholder), and
    # the first segment text is appended after it.
    assert "starts]" in first_marker
    assert "(Foreign Language)" not in first_marker
    assert "first phrase" in first_marker


def test_smtv_writer_registry_adapter_raises():
    """The 2-arg registry entry must raise so a caller that bypasses
    _write_outputs' special case fails loudly instead of dropping the
    language/work_title."""
    from core.writers import smtv_docx_writer

    with pytest.raises(RuntimeError):
        smtv_docx_writer.write([], "")


def test_smtv_output_path_uses_team_filename():
    """core.transcriber composes a fixed, recognisable team filename
    (forcing a .docx extension, not '.smtv_docx')."""
    import core.transcriber as tr

    assert tr._FMT_EXTENSIONS["smtv_docx"] == "docx"
    path = tr._smtv_output_path(os.path.join("dir", "My Show"), "ko")
    name = os.path.basename(path)
    assert name == "My Show -Transcription in Korean – Translation in English.docx"
    # Unknown language keeps the template's literal "..." placeholder.
    name2 = os.path.basename(tr._smtv_output_path("Clip", ""))
    assert name2.startswith("Clip -Transcription in ... – Translation in English")
